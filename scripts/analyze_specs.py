from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

import pypdfium2 as pdfium
from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree
from PIL import Image


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS = {"w": W_NS, "r": R_NS}


@dataclass(frozen=True)
class MediaItem:
    path: str
    bytes: int
    sha256: str
    width: int | None
    height: int | None
    mode: str | None


def safe_slug(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "_", value).strip("_")


def iter_block_items(parent: DocumentObject | _Cell) -> Iterable[Paragraph | Table]:
    parent_element = parent.element.body if isinstance(parent, DocumentObject) else parent._tc
    for child in parent_element.iterchildren():
        if child.tag == f"{{{W_NS}}}p":
            yield Paragraph(child, parent)
        elif child.tag == f"{{{W_NS}}}tbl":
            yield Table(child, parent)


def normalize_text(value: str) -> str:
    return re.sub(r"[\t\r ]+", " ", value).strip()


def paragraph_text(paragraph: Paragraph) -> str:
    chunks = paragraph._p.xpath(".//w:t/text() | .//w:tab/text()")
    return normalize_text("".join(chunks))


def xml_text(root: etree._Element, include_deleted: bool = False) -> str:
    xpath = ".//w:t/text()"
    if include_deleted:
        xpath += " | .//w:delText/text()"
    return normalize_text(" ".join(root.xpath(xpath, namespaces=NS)))


def extract_ordered_blocks(document: DocumentObject) -> list[dict[str, object]]:
    blocks: list[dict[str, object]] = []
    for item in iter_block_items(document):
        if isinstance(item, Paragraph):
            text = paragraph_text(item)
            if text:
                blocks.append(
                    {
                        "type": "paragraph",
                        "style": item.style.name if item.style else None,
                        "text": text,
                    }
                )
        else:
            rows = []
            for row in item.rows:
                rows.append([normalize_text(" ".join(cell.text.splitlines())) for cell in row.cells])
            blocks.append({"type": "table", "rows": rows})
    return blocks


def extract_headers_and_footers(document: DocumentObject) -> dict[str, list[dict[str, str]]]:
    results: dict[str, list[dict[str, str]]] = {"headers": [], "footers": []}
    seen: set[tuple[str, str]] = set()
    for index, section in enumerate(document.sections, start=1):
        for kind, container in (("headers", section.header), ("footers", section.footer)):
            text = normalize_text(" ".join(p.text for p in container.paragraphs))
            key = (kind, text)
            if text and key not in seen:
                seen.add(key)
                results[kind].append({"section": str(index), "text": text})
    return results


def read_xml(archive: zipfile.ZipFile, name: str) -> etree._Element | None:
    try:
        return etree.fromstring(archive.read(name))
    except KeyError:
        return None


def extract_comments(archive: zipfile.ZipFile) -> list[dict[str, str]]:
    root = read_xml(archive, "word/comments.xml")
    if root is None:
        return []
    comments = []
    for comment in root.xpath(".//w:comment", namespaces=NS):
        comments.append(
            {
                "id": comment.get(f"{{{W_NS}}}id", ""),
                "author": comment.get(f"{{{W_NS}}}author", ""),
                "date": comment.get(f"{{{W_NS}}}date", ""),
                "text": xml_text(comment, include_deleted=True),
            }
        )
    return comments


def extract_change_summary(archive: zipfile.ZipFile) -> dict[str, object]:
    root = read_xml(archive, "word/document.xml")
    if root is None:
        return {"insertions": 0, "deletions": 0, "inserted_text": [], "deleted_text": []}
    insertions = root.xpath(".//w:ins", namespaces=NS)
    deletions = root.xpath(".//w:del", namespaces=NS)
    return {
        "insertions": len(insertions),
        "deletions": len(deletions),
        "inserted_text": [xml_text(node) for node in insertions if xml_text(node)],
        "deleted_text": [xml_text(node, include_deleted=True) for node in deletions if xml_text(node, True)],
    }


def extract_links(archive: zipfile.ZipFile) -> list[str]:
    root = read_xml(archive, "word/_rels/document.xml.rels")
    if root is None:
        return []
    links = []
    for relationship in root:
        target = relationship.get("Target", "")
        mode = relationship.get("TargetMode", "")
        if mode == "External" and target:
            links.append(target)
    return sorted(set(links))


def extract_media(archive: zipfile.ZipFile) -> list[MediaItem]:
    media: list[MediaItem] = []
    for info in archive.infolist():
        if not info.filename.startswith("word/media/") or info.is_dir():
            continue
        payload = archive.read(info.filename)
        width = height = None
        mode = None
        try:
            from io import BytesIO

            with Image.open(BytesIO(payload)) as image:
                width, height = image.size
                mode = image.mode
        except Exception:
            pass
        media.append(
            MediaItem(
                path=info.filename,
                bytes=len(payload),
                sha256=hashlib.sha256(payload).hexdigest(),
                width=width,
                height=height,
                mode=mode,
            )
        )
    return media


def extract_pdf(pdf_path: Path, output_dir: Path, scale: float) -> tuple[list[dict[str, object]], int]:
    page_dir = output_dir / "pages"
    page_dir.mkdir(parents=True, exist_ok=True)
    for stale_preview in page_dir.glob("page-*.png"):
        stale_preview.unlink()

    rendered = pdfium.PdfDocument(str(pdf_path))
    page_count = len(rendered)
    fixed_pages = set(range(1, min(page_count, 5) + 1))
    fixed_pages.update(range(max(1, page_count - 2), page_count + 1))
    fixed_pages.update(range(1, page_count + 1, 40))
    selected_pages = set(fixed_pages)

    pages: list[dict[str, object]] = []
    for page_number in sorted(selected_pages):
        index = page_number - 1
        output_path = page_dir / f"page-{page_number}.png"
        image = rendered[index].render(scale=scale).to_pil()
        image.save(output_path, format="PNG", optimize=True)
        pages.append(
            {
                "page": page_number,
                "image": output_path.as_posix(),
                "width": image.width,
                "height": image.height,
            }
        )
    return pages, page_count


def markdown_for(report: dict[str, object]) -> str:
    lines = [f"# {report['title']}", ""]
    lines.extend(
        [
            f"- Source: `{report['source']}`",
            f"- Pages: {report['page_count']}",
            f"- Ordered content blocks: {len(report['blocks'])}",
            f"- Embedded media: {len(report['media'])}",
            f"- Comments: {len(report['comments'])}",
            f"- Tracked insertions: {report['changes']['insertions']}",
            f"- Tracked deletions: {report['changes']['deletions']}",
            "",
            "## Ordered content",
            "",
        ]
    )
    table_number = 0
    for block in report["blocks"]:
        if block["type"] == "paragraph":
            style = str(block.get("style") or "")
            text = str(block["text"])
            if style.lower().startswith("heading"):
                match = re.search(r"(\d+)", style)
                level = min(int(match.group(1)) + 1, 6) if match else 3
                lines.extend([f"{'#' * level} {text}", ""])
            else:
                lines.extend([text, ""])
        else:
            table_number += 1
            rows = block.get("rows", [])
            lines.extend([f"### Table {table_number}", ""])
            if rows:
                column_count = max(len(row) for row in rows)
                normalized = [row + [""] * (column_count - len(row)) for row in rows]
                escaped = [[str(cell).replace("|", "\\|") for cell in row] for row in normalized]
                lines.append("| " + " | ".join(escaped[0]) + " |")
                lines.append("| " + " | ".join(["---"] * column_count) + " |")
                for row in escaped[1:]:
                    lines.append("| " + " | ".join(row) + " |")
            lines.append("")

    if report["comments"]:
        lines.extend(["## Comments", ""])
        for comment in report["comments"]:
            lines.append(f"- {comment['author']}: {comment['text']}")
        lines.append("")

    if report["links"]:
        lines.extend(["## External links", ""])
        lines.extend(f"- {link}" for link in report["links"])
        lines.append("")

    lines.extend(["## Representative rendered pages", ""])
    for page in report["pages"]:
        lines.extend([f"- Page {page['page']}: `{page['image']}`", ""])
    return "\n".join(lines).rstrip() + "\n"


def analyze_document(docx_path: Path, rendered_root: Path, output_root: Path, scale: float) -> dict[str, object]:
    slug = safe_slug(docx_path.stem)
    pdf_path = rendered_root / slug / f"{slug}.pdf"
    if not pdf_path.exists():
        raise FileNotFoundError(f"Missing rendered PDF for {docx_path.name}: {pdf_path}")

    document = Document(docx_path)
    blocks = extract_ordered_blocks(document)
    headers_footers = extract_headers_and_footers(document)
    with zipfile.ZipFile(docx_path) as archive:
        comments = extract_comments(archive)
        changes = extract_change_summary(archive)
        links = extract_links(archive)
        media = [asdict(item) for item in extract_media(archive)]

    document_output = output_root / slug
    document_output.mkdir(parents=True, exist_ok=True)
    pages, page_count = extract_pdf(pdf_path, document_output, scale)

    report: dict[str, object] = {
        "title": docx_path.stem,
        "source": docx_path.as_posix(),
        "source_sha256": hashlib.sha256(docx_path.read_bytes()).hexdigest(),
        "pdf": pdf_path.as_posix(),
        "page_count": page_count,
        "blocks": blocks,
        "headers": headers_footers["headers"],
        "footers": headers_footers["footers"],
        "comments": comments,
        "changes": changes,
        "links": links,
        "media": media,
        "pages": pages,
    }
    (document_output / "report.json").write_text(json.dumps(report, indent=2, ensure_ascii=False), encoding="utf-8")
    (document_output / "content.md").write_text(markdown_for(report), encoding="utf-8")
    return report


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract and render PetSaathi DOCX specifications.")
    parser.add_argument("--docx-dir", type=Path, default=Path("DOCX"))
    parser.add_argument("--rendered-dir", type=Path, default=Path("analysis/rendered"))
    parser.add_argument("--output-dir", type=Path, default=Path("analysis/specs"))
    parser.add_argument("--scale", type=float, default=1.15, help="Representative-page render scale.")
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    reports = []
    for docx_path in sorted(args.docx_dir.glob("*.docx")):
        print(f"Analyzing {docx_path.name}", flush=True)
        reports.append(analyze_document(docx_path, args.rendered_dir, args.output_dir, args.scale))

    index = {
        "documents": [
            {
                "title": report["title"],
                "source": report["source"],
                "sha256": report["source_sha256"],
                "pages": report["page_count"],
                "blocks": len(report["blocks"]),
                "media": len(report["media"]),
                "comments": len(report["comments"]),
                "tracked_insertions": report["changes"]["insertions"],
                "tracked_deletions": report["changes"]["deletions"],
                "external_links": len(report["links"]),
            }
            for report in reports
        ]
    }
    (args.output_dir / "index.json").write_text(json.dumps(index, indent=2, ensure_ascii=False), encoding="utf-8")
    print(json.dumps(index, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
